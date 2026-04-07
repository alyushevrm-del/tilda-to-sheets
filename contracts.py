"""
Contract generators for Калий-Баскет tournament booking system.
Generates 5 documents per team: transport contract, transport appendix,
food contract, food appendix, accommodation contract.
"""

import io
import os
import re
from datetime import datetime, timedelta, date, time

from docx import Document
import openpyxl

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

MONTHS_RU = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def tpl(filename: str) -> str:
    return os.path.join(TEMPLATES_DIR, filename)


def fmt_date(d) -> str:
    """DD.MM.YYYY"""
    if isinstance(d, str):
        return d
    return d.strftime("%d.%m.%Y")


def fmt_date_month(d) -> str:
    """"DD" месяц YYYYг."""
    if isinstance(d, (datetime, date)):
        day = d.day if isinstance(d, date) else d.day
        month = d.month if isinstance(d, date) else d.month
        year = d.year if isinstance(d, date) else d.year
        return f'"{day:02d}" {MONTHS_RU[month - 1]} {year}г.'
    return str(d)


def amount_words(n: int) -> str:
    """Convert integer to Russian words (capitalised). Handles 0–999 999."""
    ones = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
            "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
            "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    # Feminine forms for thousands
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
              "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
              "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]

    def _three(num: int, feminine: bool = False) -> list[str]:
        parts = []
        h = num // 100
        rem = num % 100
        if h:
            parts.append(hundreds[h])
        if rem < 20:
            w = (ones_f if feminine else ones)[rem]
            if w:
                parts.append(w)
        else:
            t = rem // 10
            o = rem % 10
            if t:
                parts.append(tens[t])
            w = (ones_f if feminine else ones)[o]
            if w:
                parts.append(w)
        return parts

    n = int(n)
    if n == 0:
        return "ноль"

    parts = []
    millions = n // 1_000_000
    thousands = (n % 1_000_000) // 1_000
    remainder = n % 1_000

    if millions:
        mp = _three(millions)
        last2 = millions % 100
        last1 = millions % 10
        if 11 <= last2 <= 14:
            suf = "миллионов"
        elif last1 == 1:
            suf = "миллион"
        elif 2 <= last1 <= 4:
            suf = "миллиона"
        else:
            suf = "миллионов"
        parts.extend(mp)
        parts.append(suf)

    if thousands:
        tp = _three(thousands, feminine=True)
        last2 = thousands % 100
        last1 = thousands % 10
        if 11 <= last2 <= 14:
            suf = "тысяч"
        elif last1 == 1:
            suf = "тысяча"
        elif 2 <= last1 <= 4:
            suf = "тысячи"
        else:
            suf = "тысяч"
        parts.extend(tp)
        parts.append(suf)

    if remainder:
        parts.extend(_three(remainder))

    result = " ".join(p for p in parts if p)
    return result.capitalize()


def parse_arrival_dt(val) -> datetime | None:
    """Parse arrival/departure cell value to datetime."""
    if isinstance(val, datetime):
        return val
    if isinstance(val, str):
        for fmt in ("%d.%m.%Y %H:%M:%S", "%d.%m.%Y %H:%M", "%d.%m.%Y"):
            try:
                return datetime.strptime(val.strip(), fmt)
            except ValueError:
                pass
    return None


def parse_transport_str(s: str):
    """
    Parse BT/BU transport string, e.g.
    "09.04.2026 22:43 ж/д вокзал Самара, поезд 147"
    Returns (datetime | None)
    """
    if not s or "самостоятельно" in s.lower():
        return None
    m = re.search(r"(\d{2}\.\d{2}\.\d{4})\s+(\d{2}:\d{2})", s)
    if m:
        try:
            return datetime.strptime(f"{m.group(1)} {m.group(2)}", "%d.%m.%Y %H:%M")
        except ValueError:
            pass
    return None


def _replace_in_para_xml(para_el, old: str, new: str) -> bool:
    """
    Replace `old` with `new` in paragraph XML element.
    Handles text split across multiple w:t elements.
    Returns True if replacement was made.
    """
    all_t = list(para_el.iter(f"{WNS}t"))
    full = "".join(t.text or "" for t in all_t)
    if old not in full:
        return False
    replaced = full.replace(old, new)
    if all_t:
        all_t[0].text = replaced
        for t in all_t[1:]:
            t.text = ""
    return True


def replace_in_doc(doc: Document, old: str, new: str) -> None:
    """Replace text throughout the whole document (paragraphs + tables)."""
    body = doc.element.body
    for para_el in body.iter(f"{WNS}p"):
        _replace_in_para_xml(para_el, old, new)


def replace_amount_in_doc(doc: Document, cost: int, search_pattern: str) -> None:
    """
    Find paragraph matching search_pattern (regex), replace numeric amount + words.
    e.g. "72000 (Семьдесят две тысячи) рублей" → "50150 (Пятьдесят тысяч сто пятьдесят) рублей"
    """
    words = amount_words(cost)
    # Build replacement: number + space + (words) + " рублей"
    # We find the old number+words block and replace
    pattern = re.compile(r"\d[\d\s]*(?: тысяч[\w]*)?\s*\([^)]+\)\s*рублей")
    for para_el in doc.element.body.iter(f"{WNS}p"):
        all_t = list(para_el.iter(f"{WNS}t"))
        full = "".join(t.text or "" for t in all_t)
        if search_pattern not in full:
            continue
        m = pattern.search(full)
        if m:
            # Format: "50 150 (Пятьдесят тысяч сто пятьдесят) рублей"
            cost_str = f"{cost:,}".replace(",", " ")
            new_block = f"{cost_str} ({words}) рублей"
            new_full = full[: m.start()] + new_block + full[m.end() :]
            if all_t:
                all_t[0].text = new_full
                for t in all_t[1:]:
                    t.text = ""
            break


# ─────────────────────────────────────────────
# 1. Договор перевозки
# ─────────────────────────────────────────────

def generate_transport_contract(data: dict) -> bytes:
    doc = Document(tpl("transport_contract.docx"))
    arrival_dt: datetime = data["arrival_dt"]
    contract_date = arrival_dt - timedelta(days=4)

    # Blank contract number
    replace_in_doc(doc, "0000-000211", "")

    # Contract date
    replace_in_doc(doc, "20.03.2026г.", fmt_date(contract_date) + "г.")

    # Signature cell: Table[2][1] → replace blanks with contact person FIO
    t = doc.tables[0]
    cell = t.rows[2].cells[1]
    for para_el in cell._tc.iter(f"{WNS}p"):
        _replace_in_para_xml(para_el, "____________________", data["contact_person"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# 2. Приложение 1 к перевозке
# ─────────────────────────────────────────────

def generate_transport_appendix(data: dict) -> bytes:
    wb = openpyxl.load_workbook(tpl("transport_appendix.xlsx"))
    ws = wb.active

    arrival_dt: datetime = data["arrival_dt"]
    departure_dt: datetime = data["departure_dt"]
    contract_date = arrival_dt - timedelta(days=4)
    cost_transfer = int(data.get("cost_transfer") or 0)
    total_people = data["children"] + data["coaches"] + data.get("parents", 0)

    # Row 3: header
    ws["A3"] = (
        f'Приложение № 1 к договору  №  от {fmt_date_month(contract_date)}'
    )

    # Row 6: arrival date (col D=4)
    ws.cell(row=6, column=4).value = arrival_dt.date()

    # Row 8: departure date (col D=4)
    ws.cell(row=8, column=4).value = departure_dt.date()

    # Row 11: bus times
    arrival_transport_str = str(data.get("arrival_transport") or "")
    arr_train_dt = parse_transport_str(arrival_transport_str)
    if arr_train_dt:
        bus_arr_dt = arr_train_dt + timedelta(minutes=10)
        ws.cell(row=11, column=8).value = bus_arr_dt.strftime("%H:%M")
        ws.cell(row=11, column=9).value = fmt_date(bus_arr_dt.date())
    else:
        ws.cell(row=11, column=8).value = "самостоятельно"
        ws.cell(row=11, column=9).value = ""

    # Departure bus time = departure_dt time + date
    # Row 11: L12 header="время" → L11 (col 12) for time; M11:O11 merged → M11 (col 13) for date
    ws.cell(row=11, column=12).value = departure_dt.strftime("%H:%M")
    ws.cell(row=11, column=13).value = fmt_date(departure_dt.date())

    # Row 13: bus count
    ws.cell(row=13, column=6).value = 1

    # Row 14: people count
    ws.cell(row=14, column=6).value = total_people

    # Row 15: cost (override formula)
    ws.cell(row=15, column=6).value = cost_transfer

    # Row 18: cost value
    ws.cell(row=18, column=11).value = cost_transfer

    # Row 19: итого (override formula)
    ws.cell(row=19, column=11).value = cost_transfer

    # Customer info
    ws.cell(row=22, column=4).value = data["contact_person"]
    ws.cell(row=23, column=4).value = data.get("org_name", "")
    ws.cell(row=24, column=4).value = data.get("phone", "")
    ws.cell(row=25, column=4).value = data.get("email", "")

    # Payment deadline row 30
    ws.cell(row=30, column=5).value = f"до {fmt_date(arrival_dt.date())}"
    ws.cell(row=30, column=7).value = cost_transfer

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# 3. Договор питания
# ─────────────────────────────────────────────

def generate_food_contract(data: dict) -> bytes:
    doc = Document(tpl("food_contract.docx"))
    arrival_dt: datetime = data["arrival_dt"]
    departure_dt: datetime = data["departure_dt"]
    contract_date = arrival_dt - timedelta(days=4)
    cost_food = int(data.get("cost_food") or 0)

    # Blank contract number
    replace_in_doc(doc, "0000-000212", "")

    # Contract date (питание uses "г" without dot)
    replace_in_doc(doc, "20.03.2026г", fmt_date(contract_date) + "г")

    # Service period
    replace_in_doc(
        doc,
        "с 22.03.2026. по 26.03.2026",
        f"с {fmt_date(arrival_dt.date())}. по {fmt_date(departure_dt.date())}",
    )

    # Amount in words
    replace_amount_in_doc(doc, cost_food, "Стоимость оказываемых услуг")

    # Signature FIO in table row 11 (0-indexed), cell index 4
    t = doc.tables[0]
    if len(t.rows) > 11:
        cell = t.rows[11].cells[4]
        for para_el in cell._tc.iter(f"{WNS}p"):
            _replace_in_para_xml(para_el, "______________", data["contact_person"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# 4. Приложение 1 к питанию
# ─────────────────────────────────────────────

def generate_food_appendix(data: dict) -> bytes:
    wb = openpyxl.load_workbook(tpl("food_appendix.xlsx"))
    ws = wb.active

    arrival_dt: datetime = data["arrival_dt"]
    departure_dt: datetime = data["departure_dt"]
    contract_date = arrival_dt - timedelta(days=4)

    cost_food = int(data.get("cost_food") or 0)
    nights = int(data.get("nights_food") or data.get("nights_accommodation") or 1)
    children = data["children"]
    coaches = data["coaches"]
    price_food = int(data.get("price_food") or 1200)
    tournament = data.get("tournament_name", "")

    # Row 3: header
    ws["A3"] = f'Приложение № 1 к договору №  от {fmt_date_month(contract_date)}'

    # Row 9-14: counts / name
    ws.cell(row=9, column=6).value = nights
    ws.cell(row=10, column=6).value = children
    ws.cell(row=11, column=6).value = coaches
    ws.cell(row=14, column=6).value = tournament
    ws.cell(row=15, column=6).value = cost_food  # override =N37

    # ── Meal schedule (rows 26, 28, 30) ─────────────
    meal_schedule: dict = data.get("meal_schedule", {})

    arr_date_str = fmt_date(arrival_dt.date())
    dep_date_str = fmt_date(departure_dt.date())

    # Arrival day
    arr_meals = meal_schedule.get(arr_date_str, [])
    ws.cell(row=26, column=2).value = arrival_dt.date()
    ws.cell(row=26, column=8).value = ", ".join(arr_meals) if arr_meals else "ужин"

    # Middle days
    mid_start = arrival_dt.date() + timedelta(days=1)
    mid_end = departure_dt.date() - timedelta(days=1)
    if mid_start <= mid_end:
        if mid_start == mid_end:
            mid_range = fmt_date(mid_start)
        else:
            mid_range = f"{fmt_date(mid_start)}-{fmt_date(mid_end)}"
        ws.cell(row=28, column=2).value = mid_range
        ws.cell(row=28, column=8).value = "завтрак, обед, ужин"
    else:
        ws.cell(row=28, column=2).value = ""
        ws.cell(row=28, column=8).value = ""

    # Departure day
    dep_meals = meal_schedule.get(dep_date_str, [])
    ws.cell(row=30, column=2).value = departure_dt.date()
    ws.cell(row=30, column=8).value = ", ".join(dep_meals) if dep_meals else "завтрак, обед"

    # ── Cost table (rows 34-37) ──────────────────────
    children_cost = children * price_food * nights
    coaches_cost = coaches * price_food * nights

    ws.cell(row=34, column=2).value = (
        f"3-х разовое питание с {fmt_date(arrival_dt.date())}-"
        f"{fmt_date(departure_dt.date())} спортсмены"
    )
    ws.cell(row=34, column=11).value = children
    ws.cell(row=34, column=12).value = price_food
    ws.cell(row=34, column=13).value = nights
    ws.cell(row=34, column=14).value = children_cost

    ws.cell(row=35, column=2).value = (
        f"3-х разовое питание с {fmt_date(arrival_dt.date())}-"
        f"{fmt_date(departure_dt.date())} тренер"
    )
    ws.cell(row=35, column=11).value = coaches
    ws.cell(row=35, column=12).value = price_food
    ws.cell(row=35, column=13).value = nights
    ws.cell(row=35, column=14).value = coaches_cost

    # Clear dry ration row
    ws.cell(row=36, column=2).value = ""
    ws.cell(row=36, column=11).value = 0
    ws.cell(row=36, column=12).value = 0
    ws.cell(row=36, column=13).value = 0
    ws.cell(row=36, column=14).value = 0

    # Итого
    ws.cell(row=37, column=14).value = cost_food

    # ── Customer info (rows 40-43) ───────────────────
    # Row 40 col D already has "АССОЦИАЦИЯ «БК «КАЛИЙ-БАСКЕТ»" — leave as is
    ws.cell(row=41, column=4).value = data.get("org_name", "")
    ws.cell(row=42, column=4).value = data.get("phone", "")
    ws.cell(row=43, column=4).value = data.get("email", "")

    # Payment deadline row 48
    ws.cell(row=48, column=5).value = f"до {fmt_date(arrival_dt.date())}"
    ws.cell(row=48, column=7).value = cost_food

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# 5. Договор проживания
# ─────────────────────────────────────────────

def generate_accommodation_contract(data: dict) -> bytes:
    doc = Document(tpl("accommodation_contract.docx"))
    arrival_dt: datetime = data["arrival_dt"]
    departure_dt: datetime = data["departure_dt"]
    contract_date = arrival_dt - timedelta(days=4)
    cost_acc = int(data.get("cost_accommodation") or 0)
    children = data["children"]
    coaches = data["coaches"]

    # Blank contract number
    replace_in_doc(doc, "0000-000210", "")

    # Contract date (проживание: "20.03.2026г.")
    replace_in_doc(doc, "20.03.2026г.", fmt_date(contract_date) + "г.")
    # Also handle without dot at end
    replace_in_doc(doc, "20.03.20", fmt_date(contract_date)[:8])

    # Participants in preface
    old_participants = "14 спортсменов и их сопровождающий -1 тренер"
    new_participants = (
        f"{children} спортсменов и их сопровождающий - {coaches} тренер"
    )
    replace_in_doc(doc, old_participants, new_participants)

    # Service period
    replace_in_doc(
        doc,
        "с 22.03.2026г. по 26.03.2026г.",
        f"с {fmt_date(arrival_dt.date())}г. по {fmt_date(departure_dt.date())}г.",
    )

    # Check-in / check-out times (from arrival/departure datetime columns)
    checkin_time = arrival_dt.strftime("%H:%M")
    checkout_time = departure_dt.strftime("%H:%M")
    replace_in_doc(doc, "16:00", checkin_time)
    replace_in_doc(doc, "18:30", checkout_time)

    # Amount in words
    replace_amount_in_doc(doc, cost_acc, "Стоимость услуг по настоящему договору составляет")

    # Customer FIO in table signature (last table, last row, right cell)
    for t in doc.tables:
        last_row = t.rows[-1]
        for cell in last_row.cells:
            if "______________" in cell.text:
                for para_el in cell._tc.iter(f"{WNS}p"):
                    _replace_in_para_xml(
                        para_el, "______________", data["contact_person"]
                    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# Master generator
# ─────────────────────────────────────────────

def generate_all_contracts(payload: dict) -> dict:
    """
    Generate all 5 contract documents.
    Returns dict with file bytes keyed by document type.
    """
    # Parse datetimes
    payload["arrival_dt"] = parse_arrival_dt(payload.get("arrival"))
    payload["departure_dt"] = parse_arrival_dt(payload.get("departure"))

    if not payload["arrival_dt"] or not payload["departure_dt"]:
        raise ValueError(
            f"Cannot parse arrival/departure: {payload.get('arrival')!r} / {payload.get('departure')!r}"
        )

    # Extract org name (first line of "team" field, before the contact person name)
    team_str = payload.get("team", "")
    lines = [l.strip() for l in team_str.split("\n") if l.strip()]
    # The team cell often has org name + contact on same or different lines
    payload.setdefault("org_name", team_str.split(",")[0].strip() if "," in team_str else team_str)

    return {
        "transport_contract": generate_transport_contract(payload),
        "transport_appendix": generate_transport_appendix(payload),
        "food_contract": generate_food_contract(payload),
        "food_appendix": generate_food_appendix(payload),
        "accommodation_contract": generate_accommodation_contract(payload),
    }
