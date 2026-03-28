import os
import io
import json
import logging
import re
import base64
from datetime import datetime, timedelta

from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse

import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

SPREADSHEET_ID = os.environ["SPREADSHEET_ID"]
FIRST_DATA_ROW = 5
SUMMARY_COL = "BY"
LINK_COL = "BZ"

app = FastAPI(title="Tilda to Google Sheets webhook")

BUSES = [
    (18, 16750),
    (20, 20500),
    (25, 20250),
    (35, 29250),
    (49, 37500),
    (58, 45000),
]

MEAL_MAP = {
    "завтрак": 0,
    "обед": 1,
    "полдник": 2,
    "ужин": 3,
    "второй ужин": 4,
}
STANDARD_MEAL_OFFSETS = [0, 1, 3]  # з, о, у

TEMPLATE_B64 = "UEsDBBQABgAIAAAAIQDfpNJsWgEAACAFAAATAAgCW0NvbnRlbnRfVHlwZXNdLnhtbCCiBAIooAACAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC0lMtuwjAQRfeV+g+Rt1Vi6KKqKgKLPpYtUukHGHsCVv2Sx7z+vhMCUVUBkQpsIiUz994zVsaD0dqabAkRtXcl6xc9loGTXmk3K9nX5C1/ZBkm4ZQw3kHJNoBsNLy9GUw2ATAjtcOSzVMKT5yjnIMVWPgAjiqVj1Ykeo0zHoT8FjPg973eA5feJXApT7UHGw5eoBILk7LXNX1uSCIYZNlz01hnlUyEYLQUiep86dSflHyXUJBy24NzHfCOGhg/mFBXjgfsdB90NFEryMYipndhqYuvfFRcebmwpCxO2xzg9FWlJbT62i1ELwGRztyaoq1Yod2e/ygHpo0BvDxF49sdDymR4BoAO+dOhBVMP69G8cu8E6Si3ImYGrg8RmvdCZFoA6F59s/m2NqciqTOcfQBaaPjP8ber2ytzmngADHp039dm0jWZ88H9W2gQB3I5tv7bfgDAAD//wMAUEsDBBQABgAIAAAAIQAekRq37wAAAE4CAAALAAgCX3JlbHMvLnJlbHMgogQCKKAAAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArJLBasMwDEDvg/2D0b1R2sEYo04vY9DbGNkHCFtJTBPb2GrX/v082NgCXelhR8vS05PQenOcRnXglF3wGpZVDYq9Cdb5XsNb+7x4AJWFvKUxeNZw4gyb5vZm/cojSSnKg4tZFYrPGgaR+IiYzcAT5SpE9uWnC2kiKc/UYySzo55xVdf3mH4zoJkx1dZqSFt7B6o9Rb6GHbrOGX4KZj+xlzMtkI/C3rJdxFTqk7gyjWop9SwabDAvJZyRYqwKGvC80ep6o7+nxYmFLAmhCYkv+3xmXBJa/ueK5hk/Nu8hWbRf4W8bnF1B8wEAAP//AwBQSwMEFAAGAAgAAAAhANZks1H0AAAAMQMAABwACAF3b3JkL19yZWxzL2RvY3VtZW50LnhtbC5yZWxzIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArJLLasMwEEX3hf6DmH0tO31QQuRsSiHb1v0ARR4/qCwJzfThv69ISevQYLrwcq6Yc8+ANtvPwYp3jNR7p6DIchDojK971yp4qR6v7kEQa1dr6x0qGJFgW15ebJ7Qak5L1PWBRKI4UtAxh7WUZDocNGU+oEsvjY+D5jTGVgZtXnWLcpXndzJOGVCeMMWuVhB39TWIagz4H7Zvmt7ggzdvAzo+UyE/cP+MzOk4SlgdW2QFkzBLRJDnRVZLitAfi2Myp1AsqsCjxanAYZ6rv12yntMu/rYfxu+wmHO4WdKh8Y4rvbcTj5/oKCFPPnr5BQAA//8DAFBLAwQUAAYACAAAACEAMs26uyYHAAA8nAAAEQAAAHdvcmQvZG9jdW1lbnQueG1s7J1Lb9tGEMfvBfodCJ0Tk5Qs6oHYQWI3aYC2MJK0Pa/IlcSG5BLkyrJ6StIURRGgvRRFUKApWuTWi9E6qZuHC+QTUF+hn6SzfOhhSgop17Ykz8Xk7nL/nJ2Znd/aCJgrV/dsS9qlnm8yZ6OgrikFiTo6M0yntVH49O6Ny9WC5HPiGMRiDt0o9KhfuLr5/ntXunWD6R2bOlwCCcevd119o9Dm3K3Lsq+3qU38NdvUPeazJl/TmS2zZtPUqdxlniEXFVUJ71yP6dT34X1bxNklfiGW0/eyqRke6cJkIbgu623icbo31FBzi5TlmlxNCxXnEIIVFtW0VCm3lCYLq1JC63MJgVUppfJ8ShMWp82nVEwrVeZTKqWVqvMppdLJTic4c6kDg03m2YRD02vJNvHuddzLIOwSbjZMy+Q90FS0RIaYzr05LIJZAwW7ZORWqMg2M6hVMhIVtlHoeE49nn95MF+YXo/mx5fBDGpley28ribTPW75PJnrZfFdNH07Liyh12SPWuBH5vht0x1UB3teNRhsJyK7sxywa1vJc11XzbjVppW27SgMQ8Es5sexs63I8tmKqpIhmkJiMCOLCePvTCyxIYOHL57LNSPOVTMWn0SgmBLQdJoRFolGNdaQ9eHuFjpmxm2V6ERRETrm0LFqxhp43JgRAd/gRjuXSjHxqyzmEk7axB8kulCk+YwqD+R69oiP3NbJNsJNj3XcoZp5MrVbw5LYFYeTHFrxhhrd5P7JjLnTJi5USluv32o5zCMNCyyC7SFBhkthBMRPSBRxCW/pXtgvYi2JGlPYhFNVgxk9cXWlbh1OZcbtjYKibG2XK9XrhaRrx5vQuU2bpGPx9MjOSFeovOOJyxc6jO8SqOg6VEfqFWTR60WD3g3mcB8eIL5uQpTumjb1pU9oV7rNbOII6fY1x588ovvpbhCXY3V5YII3cUFnZEa3Hp5c675LdIiU61Gferu0sBn8GvwTHPYfBEfBSyn4K9gPnsPPF8F+/7v+tzDwtRS8Cfal/sP+V/37cHvYvy8JVR5pjy1MLKuoVmpF7QyXtfn29+BnMPsweA2mv4F1/AEtMDR4LgWHUnAAS3nYfzy2gmD/7atpizif6DS8uJly6PWyek3bPm1LwIIcbn0Jzddh86D/+JIU/BkcgVOPgoNFc+uUpJ+awYsU/ItqSzoVg98gu0QSPu/fX4r8qs8skWezo8cc+Cx4Ejy9FBdDUc7Fbn0BTeHVw/73l0SVF3saPAyVErpfwe0j6ILij5slgy2C86GzG1Z8iZ3fsO7wngXHoej4QUpRjkP/59AHv5go4lW850IGkQ5ng+GPGLuXzFLWr4WPNU3P57cZzFJF0yJxazi4xayOLf64lYwnHeEjDvvwOnGMQeuzqKXGyxq1+qZnGuK2BVfQiGytrmuRfWO9xWqtOKFbrWq1/N2RGcnb+djOyXAyvHs8Ebge/YwXpsdeFysZ+N3YI8mr4+dW70A6XlZFup7jsTRd4/999ONYnYn2UxiRGVEMMw/DuDhhDJ6tBU/WgqdrcwQzrAAYzAUK5g9TjgsY3BUI7k/idCfCG7yEn2/6j4O/jx39MMwrEOZf4HfzBxDkA7h+E57y4Rd4cdB/AME+mh5icYk0BisR69BKlUqlOjNmp3EEy/Da0VQZfzxMlbjrPFNlQZJkND3UOfZ4rkPXSgbuFGoh+mkuP/0fLpm6HTWtvL5dPLvtmPpzAmZYrgxDZF0EZBURWYgsRBYiC5GFyFoOZJUQWYgsRBYiC5GFyFoOZK0jshBZiCxEFiILkbUcyCojshBZiCxEFiILkbUcyNIQWYgsRBYiC5GFyFoOZFUQWYgsRBYiC5GFyFoOZFURWYgsRBYiC5GFyFoOZNUQWYgsRBYiC5GFyFoOZKkKMguZhcxCZiGzkFlLwiz8/gUyC5mFzFqJDENmXQhm4QcwkFnILGTWSmTYFGbF30iftfDTYNax1464dnxkgfF0ZpkPJJr0XYvZHn1XBi+Hm7Pt9FzMvcB+ylUR0U8n8FMMR/TTRconJOwSE3bSZziwIiJhsSIiYTGfkLDv9jcSdjZhJ301BCsiEhYrIhIW82mxCRv/Z8OzFn4ahD322hHXjo8gYbuCsJM+crJiHj0FmGbI7VX1U67ih346gZ+mwxT9tLL5hDBdYphO+vwKwhQ3NRY/hCnmE8IUYZoDppM+DIMwxU2NxQ9hivl0bjAVl4aVWv8HW0pJy/6PocdHwvWPwHaBiDdxQWdkRjf9D9i7dZ/qfGdg1rtcH3qzdefLOKPVmhImXRvutWqpGnnUbX1MhCJnrnimtB5qma02hKpaVkSrwThn9nDUok0YVCuKGspRYlB4bUUJzWgyxkearQ4Pm0r0Np1Zwl++S3QaPRN2G0y/6ZmG0DYdumNyHYwsaeEkOVl2eNtgRi+8gSkdGzJj8z8AAAD//wMAUEsDBBQABgAIAAAAIQCGOuCnSwYAAL4aAAAVAAAAd29yZC90aGVtZS90aGVtZTEueG1s7Fndihs3FL4v9B3E3Dv+m/HPEm+wx3bTZjcJ2U1KLuUZeUZZzchI8m5MCITNZaGlNC29aKDtTS9K24Wk9CZ5B+cZtk1pU8grVKPx2CNbZpNmA0uJDR79fOfo0znSpxnP+Qu3IwL2EeOYxi2rfK5kARR71Mdx0LKu7/YLDQtwAWMfEhqjljVB3Lqw+f575+GGCFGEgLSP+QZsWaEQo41ikXuyGfJzdIRi2TekLIJCVllQ9Bk8kH4jUqyUSrViBHFsgRhG0u30++mv0yfTI3BlOMQesjYz/z0if2LBkwaPsJ3EO8qMvnt2OD2aPp0+nh49uyfLT+X1M2Xr75WTC59wlzCwD0nLkkP79GAX3RYWIJAL2dGySupjFTfPF+dGRKyxzdn11WdmNzPw9yrKjgWDuaFtO3atPfevAESs4nr1Xq1Xm/tTAOh5cuYplzzW6TQ7XWeGzYHSosF3t96tljV8zn91Bd92kq+GV6C0aK/g+313EcMcKC06hpjUK66t4RUoLdZW8PVSu2vXNbwChQTHeyvoklOrutls55AhJReN8KZj9+uVGXyBKuZWW2ofi1ddexG8RVlfGqhkQ4FjICYjNISetHMhwQOGwRYOQrkQRzCmXDaXKqV+qSp/k6+tSipCcAPBnHXa5PGVpoQf4B7DI9GyPpJerRzk5ZOfXj55BI4PHx8f/nZ8//7x4S8Gq4swDvJWL374/J+H98Dfj7598eBLM57n8X/8/MnvT78wA0Ue+Pyroz8fHz3/+tO/fnxggLcZHOThuzhCHFxGB+AajeTEDAOgAXs9i90Q4rxFOw44jGFiY0D3RKihL08ggQZcB+kRvMGkbJiAH4xvaYR3QjYW2AC8FEYacJtS0qHMOKdLyVj5KIzjwDw4G+dx1yDcN43tLuW3Nx7J9Y9NLt0QaTSvEplyGKAYCZD00T2EDGY3Mdbiuo09RjkdCnATgw7ExpDs4oG2mhZGF3Ek8zIxEZT51mKzfQN0KDG576J9HSl3BSQml4hoYfwAjgWMjIxhRPLILShCE8mdCfO0gHMhMx0gQkHPR5ybbK6wiUb3kpQXc9q3ySTSkUzgPRNyC1KaR3bpnhvCaGTkjOMwj/2Q78klCsFVKowkqL5DkrrMA4zXpvsGRlq6T97b16WymhdI0jNmpi2BqL4fJ2QIkXJeXNLzCMcnivuSrDtvV9alkD7/5qFZd8+koLcZNu6oZRlfh1sWb5cyH5997e7CcXwVye1igL6T7nfS/b+X7nX7+fQFe6HR6qY+u3VXbqJXvo8fYkJ2xISgLa7Unsvp+n3ZqCrKyfwxYhTK4mx4DRcwqMqAUfExFuFOCEdy2LIaIeAz1wEHI8rleaGajb6TDjKOtqmftpbL2ZOrNIBi0S7Pm6xdnk4iba3VF49oc/eqFqhH64xAYvs6JHKD6SSqBhL1rPEEEmpmp8KiaWDRSNyvZaEus6zI/Qhg8j+IY6eM5PqDBPlJnlL7LLunnul1wdSnXTFMr5lwPZ1MayRyy00nkVuGIfTRcvMp57q5SKlGLwnFKo16423kOhGVJW0gsV4DB3LPVR3pxoOjljWUd4qyGI2kP57oKCRB3LI8MQv0f1GWEeOiC3mYwlRXOv8IC8QAwZFc6/k0kHjBrVypJ3M8o+SapbMXOXXJJxkNh8gTa1oWVdmXOjH2viE4qdCxJL0T+gdgQMbsGpSBcurlJIA+5mIeTR+z3OJeRHFJrmZbUftHbbFFIRmFcHai5MU8havynE5uHorp8qz0+mwygyBJ0hufuicbJR050VxzgCSnplk/3t4hn2O10H2NVSrdy1rXzLRu3Snx5gdCjtpiMI1awthAbdGqUzvFG4LccPOlue6MOO3TYHnVJgdEdp+paiuvMujgllz5XXn7OiaCK6rotnxmcLM/nVMlUK2ZutwWYMxwy7pTctq2W3HcQqnh9Ap21S4VGk67Wmg7TrXcc8qlbqdyVwZFhFHZScfuy+cbMpm9rFHtKy9souy2+5xHoyJV72GKyli9sClXtBc26XsasJv0WwDLyNypVfrNarNTKzSr7X7B7nYahaZb6xS6Nbfe7Xddp9Hs37XAvgLb7apr13qNQq3sugW7VkroN5qFul2ptO16u9Gz23dnsZYzz65ZeBWvzX8BAAD//wMAUEsDBBQABgAIAAAAIQAgGuPiOwQAADkMAAARAAAAd29yZC9zZXR0aW5ncy54bWy0Vt9v2zYQfh+w/8HQuyPJPyRbq1PEdrykiNchdrFnSqIsIqRIkJQdt9j/viMlWk6TFUm3vNjUfXffHY/HO374+Mhob4+lIryaeeFF4PVwlfGcVLuZ92W76k+8ntKoyhHlFZ55R6y8j5e//vLhkCisNaipHlBUKmHZzCu1Fonvq6zEDKkLLnAFYMElQxo+5c5nSD7Uop9xJpAmKaFEH/1BEEReS8NnXi2rpKXoM5JJrnihjUnCi4JkuP1zFvI1fhuTJc9qhittPfoSU4iBV6okQjk29rNsAJaOZP+jTewZdXqHMHjFdg9c5ieL14RnDITkGVYKDohRFyCpOsejZ0Qn3xfgu92ipQLzMLCr88jHbyMYPCOIMvz4No5Jy+GD5TkPyd/GE514SJfYMPq5YM4IVK7z8k0sA5dX39gijUqkTlVkGPHbghqf6I6sy5Gir6maBrojqUSyuZNtybAsud1VXKKUQjhQOj04/Z6NzvxCEs2fXeJHKzd58C6hR3zlnPUOicAyg4sCDWY48nwDQHnyYqORBopECUyp7TgZxQg8HpKdRAx6hZNYmxwXqKZ6i9KN5gKU9gg2FgeTBs5KJFGmsdwIlAHbgldacur0cv4H1wvoOxKuRWthu1C32jQdDSwqxGCrT7rUmufYRFZL8vozMQbWezg+d/m9Iw4dWJIcb02KN/pI8QqC35Cv+KrKP9VKE2C0veo/RPCjAHBlPH+GotgeBV5hpGtI0zs5syexokSsiZRc3lY51Ma7OSNFgSU4IFBraygfIvnB5vkGoxwG3zv5rRX+C5ThTg63UJYPc641ZzdHUUKu/4eT9M/LF8Z3rtzinnPtVIPgehpfT9riM2iHDMJ4OoheQqJhHMftrXqKzMfhVbR8CVksx/Fk/hLy7xFcL4JhZP34pz2wxAzSP6VbmYvQY43FArFUEtRbm1HrG41UPsxJ5fAUQz/D58imTh3Y7zeAYojSFRyJA4JGnhMllriwa7pGctfxthryRSl0pU8nLtPlsPxd8lo06EEi0RS4UwlHo9aSVPqOMCdXdbpxVhV04DOorvLPe2nz1KXnkGgoGNso7pAtPKsr6/79l7YwqdyYosJrJERTm+kunHmU7EodmnLS8JXDi8x+pLtBiw0sNmgw+4EyszPQbhedbOBkZ3pDJxt2spGTjTrZ2MnGnSxyssjISuhGEkbDA1wTtzTyglPKDzi/6fBnoiYJqkQCL5vJAeXFG0E7SlRvn+BHmEs4JxoeuoLkDMGjJAyai9FqU3TktX6iazCjLJ4ymBHuLugTY1vi38ViJlpGoBw3R5Z2g+qiCZwSBU1FwEzTXDrsN4uFoyTn2a0Zu6NGHq1W18H4atXAYzsLte07cO73uJgjhfMWc6bjxvTbKJqGcRyu+leTYdgfLYbT/jS6mvfjKQjiVTSIl8u/20vq3vyX/wAAAP//AwBQSwMEFAAGAAgAAAAhABrVRaOBAQAA7gIAABEACAFkb2NQcm9wcy9jb3JlLnhtbCCiBAEooAABAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHySXU+DMBSG7038D6T3rAV0LoSx+JFducQoRuNdbc+2Olqathvbv7fAxsQY73p6nvPk8JZstpdlsANjRaWmKBoRFIBiFRdqNUWvxTycoMA6qjgtKwVTdACLZvnlRcZ0yioDT6bSYJwAG3iTsinTU7R2TqcYW7YGSe3IE8o3l5WR1PnSrLCmbENXgGNCxliCo5w6ihthqHsjOio565V6a8pWwBmGEiQoZ3E0ivCZdWCk/XOg7fwgpXAHDX+ip2ZP763owbquR3XSon7/CL8vHl/aTw2FarJigPKMs9QJV0Ke4fPRn+z28wuY6677wp+ZAeoqk0uqfDImaYHTZRP3Bg51Zbj1o4PKYxwsM0I7/4ideHDh6ZJat/CvuhTA7w75bSm2dg274JnateCt8BfSTBnYiebHyJOW6MvsmHK3HvDAp5N2WZ46b8n9QzFHeUzi6zAiIbkpojglVykhH82Gg/mzUB4X+N84DkkSxpOCjNMoGRpPgi6k4R+afwMAAP//AwBQSwMEFAAGAAgAAAAhAGnvrWgAAgAAegYAABIAAAB3b3JkL2ZvbnRUYWJsZS54bWy8k92OmzAQhe8r9R2Q7zcY8rMpWrLaphupUtWL1fYBHGOCVWwjjxOSt+/YEBptFGmp1OXCmDOej5nD8PB4VHV0EBak0TlJJpREQnNTSL3Lya/Xzd2SROCYLlhttMjJSQB5XH3+9NBmpdEOIszXkCmek8q5Jotj4JVQDCamERqDpbGKOXy0u1gx+3vf3HGjGubkVtbSneKU0gXpMfY9FFOWkotvhu+V0C7kx1bUSDQaKtnAmda+h9YaWzTWcAGAPau64ykm9YBJZlcgJbk1YEo3wWb6igIK0xMadqr+C5iPA6RXgAUXx3GMZc+IMfOSI4txnMXAkcUF59+KuQBA4YpqFCU9+xr7XOZYxaC6JIpxRc0H3El5jxTPvu+0sWxbIwm/eoQfLgpgv2L//ha24hh03wJZ9b9C1GaaKcxcs1purQyBhmkDIsHYgdU5wR42dE59Lymd0alfSewP8opZEB4SDq7XnVwyJevTWYVWAnSBRjpenvUDs9JX3YVA7jCwhy3NyXNCafq02ZBOSZBMKX1aD0qKRXXXl16ZDgr1Cg+c8Jh0WTxwhjP4zrhz4MqJV6kERD9FG70YxfQNR1K6QCfm6Id3ZnrDke5Nbx2xgTvKEaw5fX7ryP1y/vUjHOlnI/ohd5W7OSF+Lj5sQrwfVxOSzu7/jx/9BlZ/AAAA//8DAFBLAwQUAAYACAAAACEA7wopTk4BAAB+AwAAFAAAAHdvcmQvd2ViU2V0dGluZ3MueG1snNNfa8IwEADw98G+Q8m7psoUKVZhDMdexmDbB4jp1YYluZKLq+7T79qpc/hi95L/9+MuIfPlztnkEwIZ9LkYDVORgNdYGL/JxfvbajATCUXlC2XRQy72QGK5uL2ZN1kD61eIkU9SwoqnzOlcVDHWmZSkK3CKhliD580Sg1ORp2EjnQof23qg0dUqmrWxJu7lOE2n4sCEaxQsS6PhAfXWgY9dvAxgWURPlanpqDXXaA2Gog6ogYjrcfbHc8r4EzO6u4Cc0QEJyzjkYg4ZdRSHj9Ju5OwvMOkHjC+AqYZdP2N2MCRHnjum6OdMT44pzpz/JXMGUBGLqpcyPt6rbGNVVJWi6lyEfklNTtvetXfkdPa08RjU2rLEr57wwyUd3LZcf9t1Q9h1620JYsEfAutonPmCFYb7gA1BkO2yshabl+dHnsg/v2bxDQAA//8DAFBLAwQUAAYACAAAACEAJu75n/oLAABpdQAADwAAAHdvcmQvc3R5bGVzLnhtbLyd23LbOBKG77dq34Glq90LR5YPcpIaZ8p2krVr48QT2ZNriIQsrElCC5I+zNMvAJIS5SYoNtjrm8QS1R9A/PibaB6k339/TuLgkatMyPR0NHm3Pwp4GspIpPeno7vbr3vvR0GWszRisUz56eiFZ6PfP/37b09f87xl5hngQak2cckPB0t83z1cTzOwiVPWPZOrniqNy6kSliuX6r7ccLUQ7HaC2WyYrmYi1jkL+OD/f3pqMKoPhS5WIiQf5ZhkfA0t/FjxWNNlGm2FKuspj31oT1JFa2UDHmW6Z1O4pKXMJGuMZMjAEpEqGQmF/k7vTNVjyxKh0/27V9JvAEc4wAHADAN+TOO8b5ijHVkkyMiHGe65oiowfHrTAOQRXm0RFEO6nEdm1iWsyXLlk0ix3XqeI17ScwYJeHHq/tUKjaPNUmrHmjhAgs2/+r9N//ZP/mzfd/swuiT9kIkw898wYo4z8xLdaOql9Ur+99XmeZZ8PSRZaEQt7qDupVE6AYvz9JMjPQWzrL8LBOsdePS/NG6JczyxtvnIhKjsWkx+0tvfGTx6ejgoH7nwvRg672Ypff1e6rY+3nX7MnpiKd7dzPz1lxzT0dM7c3OTOC42rHy/8burl6/sg2vWChsO2yRc23zyXTfQGNhssrB8Yf6xc/CDD4rclk1YgHl/2vsGIy4dr/OBbMyJemtfPFNhg88muV6w+nItqXfvLu6UUIqnXZORx9sm/rNGU/EpYginjY+mC5FxH8teXqX8Wjz/h9fbeqo3ghlkeq/D0+mdhbEWfTlOeQrk4j01pQZTb6bgNh8uhCbxm34f2vYpFKiLX7JmcnGweQ1wnYfhTgwEVljb9uZxat9t59CNXT4Vg0dvVVDx2/V0PStGjp5q4bev1VDFvP/bEikkU789vOwGUDdxXG4Ec1xmA3NcXgJzXFYBc1xOAHNcUx0NMcxj9EcxzRFcHIZumZhY7IfOmZ7N3f3McKPu/uQ4MfdfQTw4+5O+H7c3fndj7s7nftxd2dvP+7uZI3nlkut4ErbLM0Hu2whZZ7KnAc5fx5OY6lm2RKVhmcOelyR7CQBpsxs1YF4MC1k9vXuGWJN6n88z02lF8hFsBD3heLZ4I7z9JHHcsUDFkWaRwhUPC+UY0R85rTiC654GnLKiU0HNZVgkBbJnGBurtg9GYunEfHw1USSpLCe0Lp+XhqTCIJJnbBQyeFdk4wsP3wT2fCxMpDgvIhjTsT6TjPFLGt4bWAxw0sDixleGVjM8MKgoRnVEFU0opGqaEQDVtGIxq2cn1TjVtGIxq2iEY1bRRs+brcij22Kb646Jv3P3V3E0lxUGNyPmbhPmV4ADD/cVOdMgxum2L1iq2Vgzkq3Y5v7jG3nXEYvwS3FMW1NolrX2ylyofdapMXwAd2iUZlrzSOy15pHZLA1b7jFrvUy2SzQLmnqmVkxz1tNa0m9TDtjcVEuaIe7jeXDZ9jGAF+Fyshs0I4lmMHfzXLWyEmR+Ta9HN6xDWu4rV5nJdLuVUiCXsYyfKBJw5cvK650WfYwmPRVxrF84hEdcZYrWc61puUPrCS9LP8lWS1ZJmyttIXof6ivb0cIrtlq8A7dxEykNLp92UuYiAO6FcTl7fW34FauTJlpBoYGeC7zXCZkzOpM4D9+8fk/aTp4povg9IVob8+ITg9Z2IUgOMiUJBkRkfQyU6SC5Bhqef/mL3PJVERDu1G8vAMo50TEGUtW5aKDwFs6Lz7p/EOwGrK8P5kS5rwQlaluSWCN04ZZMf8PD4enuu8yIDkz9KPI7flHu9S10XS44cuELdzwJYJVUx8ezPwl2Nkt3PCd3cJR7exFzLJMOC+hevOodrfmUe/v8OKv4slYqkUR0w1gDSQbwRpINoQyLpI0o9xjyyPcYcuj3l/CKWN5BKfkLO9fSkRkYlgYlRIWRiWDhVFpYGGkAgy/Q6cBG36bTgM2/F6dEka0BGjAqOYZ6eGf6CpPA0Y1zyyMap5ZGNU8szCqeXb4OeCLhV4E0x1iGkiqOddA0h1o0pwnK6mYeiFCfon5PSM4QVrSbpRcmEdDZFrexE2ANOeoY8LFdomjEvkXn5N1zbAo+0VwRpTFsZRE59Y2BxwbuX3v2q4w+yTH4C7cxCzkSxlHXDn2yR2r6+VZ+VjG6+7bbvQ67flN3C/zYLZcn+1vYqb7OyPrgn0rbHeDbWM+rZ9naQu75pEokrqj8GGK6WH/YDujt4KPdgdvVhJbkcc9I2Gb092Rm1XyVuRJz0jY5vuekdanW5FdfvjM1EPrRDjpmj/rGs8x+U66ZtE6uLXZrom0jmybgidds2jLKsFZGJqrBVCdfp5xx/czjzse4yI3BWMnN6W3r9yILoP95I/CHNkxSdO2t757AuR9u4julTn/KGR53n7rglP/h7qu9MIpzXjQyjnsf+FqK8u4x7F3unEjeucdN6J3AnIjemUiZzgqJbkpvXOTG9E7SbkR6GwFjwi4bAXjcdkKxvtkK0jxyVYDVgFuRO/lgBuBNipEoI06YKXgRqCMCsK9jAopaKNCBNqoEIE2KlyA4YwK43FGhfE+RoUUH6NCCtqoEIE2KkSgjQoRaKNCBNqonmt7Z7iXUSEFbVSIQBsVItBGtevFAUaF8Tijwngfo0KKj1EhBW1UiEAbFSLQRoUItFEhAm1UiEAZFYR7GRVS0EaFCLRRIQJt1PJRQ3+jwnicUWG8j1EhxceokII2KkSgjQoRaKNCBNqoEIE2KkSgjArCvYwKKWijQgTaqBCBNqq9WDjAqDAeZ1QY72NUSPExKqSgjQoRaKNCBNqoEIE2KkSgjQoRKKOCcC+jQgraqBCBNipEdM3P6hKl6zb7Cf6sp/OO/f6XrqpO/Ww+yt1EHfZH1b1ys/o/i3Au5UPQ+uDhoa03+kHEPBbSnqJ2XFZvcu0tEagLnz8uup/wadIHfulS9SyEvWYK4Ed9I8E5laOuKd+MBEXeUddMb0aCVedRV/ZtRoLD4FFX0rW+rG9K0YcjENyVZhrBE0d4V7ZuhMMh7srRjUA4wl2ZuREIB7grHzcCjwOTnF9HH/ccp+n6/lJA6JqODcKJm9A1LaFWdTqGxugrmpvQVz03oa+MbgJKTycGL6wbhVbYjfKTGtoMK7W/Ud0ErNSQ4CU1wPhLDVHeUkOUn9QwMWKlhgSs1P7J2U3wkhpg/KWGKG+pIcpPangow0oNCVipIQEr9cADshPjLzVEeUsNUX5Sw8UdVmpIwEoNCVipIcFLaoDxlxqivKWGKD+pQZWMlhoSsFJDAlZqSPCSGmD8pYYob6khqktqexZlS2qUwo1w3CKsEYg7IDcCccm5EehRLTWiPaulBsGzWoJa1ZrjqqWmaG5CX/XchL4yugkoPZ0YvLBuFFphN8pPamy11Ca1v1HdBKzUuGrJKTWuWuqUGlctdUqNq5bcUuOqpTapcdVSm9T+ydlN8JIaVy11So2rljqlxlVLbqlx1VKb1LhqqU1qXLXUJvXAA7IT4y81rlrqlBpXLbmlxlVLbVLjqqU2qXHVUpvUuGrJKTWuWuqUGlctdUqNq5bcUuOqpTapcdVSm9S4aqlNaly15JQaVy11So2rljqlxlVL1zpEEHwF1CxhKg/ovi/ukmXLnA3/csK7VPFMxo88Cmh39RtqL8dPWz9/Zdj2t/n053M9ZuYb0BuPK0XlN8BWQPvBK01i9hesTCeC6rfAqh+usn2trtSWjdkY2Eq41M2E1ddWuVrZB804vpHWNruZafWnq7HbDEz5ua1h6exlbmZ2Vw8njoEoPeHq14fK5Ls6prsxj8ufRNN/XKWRBjxVPwdWdjB6ZiVKb7/gcXzNyk/LlfujMV/k5dbJvv1Kglfb5+W36znjlU3DTsB4uzPly+pn2RzDXH7ffnV/gGuoD1qG2t6oMnSUe8i/6cUh6EWZ/cxpo2r4mMb+MCa1m1k1CWH/6se+VSaMsPa9/f2Lz8cn78/LLa6fw2v+GN7R+kX7j+Ftz6FzqSKurOvLOWJbNV87Xe3oX/oQZP/QbfL1D9fpvL4hr2eQV+x6dnlF13PPK1ikeqT55bDwP/3CSxush7/NFfVf2af/AQAA//8DAFBLAwQUAAYACAAAACEA0DmYH3cBAADKAgAAEAAIAWRvY1Byb3BzL2FwcC54bWwgogQBKKAAAQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACcUstOwzAQvCPxD1HurdNSEKCtK9QKceAlNZSzZW8SC8e2bIPav2dDaAjiRk47s97xzMaw2rcm+8AQtbPLfDYt8gytdErbepm/lLeTyzyLSVgljLO4zA8Y8xU/PYHn4DyGpDFmJGHjMm9S8teMRdlgK+KU2pY6lQutSARDzVxVaYkbJ99btInNi+KC4T6hVagmfhDMe8Xrj/RfUeVk5y/uyoMnPQ4ltt6IhPyxmzRT5VILbGChdEmYUrfIF0QPAJ5FjZHPgPUFvLqgIj+fA+srWDciCJlogXx+dQlshOHGe6OlSLRa/qBlcNFVKXv68pt188DGR4AybFG+B50OvAA2hnCvLd1P9/YFGQuiDsI33+4GBFspDK4pPa+EiQjsh4C1a72wJMeGivTe4osv3aZbxPfIb3IU8lWnZuuFJAtni6tx3FEHtsSiIv+DhYGAO/ojwXT6NGtrVMczfxvdAnf90+Szi2lB39fGjhzlHt4M/wQAAP//AwBQSwECLQAUAAYACAAAACEA36TSbFoBAAAgBQAAEwAAAAAAAAAAAAAAAAAAAAAAW0NvbnRlbnRfVHlwZXNdLnhtbFBLAQItABQABgAIAAAAIQAekRq37wAAAE4CAAALAAAAAAAAAAAAAAAAAJMDAABfcmVscy8ucmVsc1BLAQItABQABgAIAAAAIQDWZLNR9AAAADEDAAAcAAAAAAAAAAAAAAAAALMGAAB3b3JkL19yZWxzL2RvY3VtZW50LnhtbC5yZWxzUEsBAi0AFAAGAAgAAAAhADLNursmBwAAPJwAABEAAAAAAAAAAAAAAAAA6QgAAHdvcmQvZG9jdW1lbnQueG1sUEsBAi0AFAAGAAgAAAAhAIY64KdLBgAAvhoAABUAAAAAAAAAAAAAAAAAPhAAAHdvcmQvdGhlbWUvdGhlbWUxLnhtbFBLAQItABQABgAIAAAAIQAgGuPiOwQAADkMAAARAAAAAAAAAAAAAAAAALwWAAB3b3JkL3NldHRpbmdzLnhtbFBLAQItABQABgAIAAAAIQAa1UWjgQEAAO4CAAARAAAAAAAAAAAAAAAAACYbAABkb2NQcm9wcy9jb3JlLnhtbFBLAQItABQABgAIAAAAIQBp761oAAIAAHoGAAASAAAAAAAAAAAAAAAAAN4dAAB3b3JkL2ZvbnRUYWJsZS54bWxQSwECLQAUAAYACAAAACEA7wopTk4BAAB+AwAAFAAAAAAAAAAAAAAAAAAOIAAAd29yZC93ZWJTZXR0aW5ncy54bWxQSwECLQAUAAYACAAAACEAJu75n/oLAABpdQAADwAAAAAAAAAAAAAAAACOIQAAd29yZC9zdHlsZXMueG1sUEsBAi0AFAAGAAgAAAAhANA5mB93AQAAygIAABAAAAAAAAAAAAAAAAAAtS0AAGRvY1Byb3BzL2FwcC54bWxQSwUGAAAAAAsACwDBAgAAYjAAAAAA"



def calculate_transfer_cost(total_people: int) -> int:
    if total_people <= 0:
        return 0
    INF = float("inf")
    dp = [INF] * (total_people + 1)
    dp[0] = 0
    for n in range(1, total_people + 1):
        for cap, cost in BUSES:
            if cap >= n:
                dp[n] = min(dp[n], cost)
            elif n - cap >= 0 and dp[n - cap] < INF:
                dp[n] = min(dp[n], dp[n - cap] + cost)
    return dp[total_people] if dp[total_people] < INF else 0


def normalize_phone(phone: str) -> str:
    if phone.startswith("+7"):
        return "8" + phone[2:]
    return phone


def col_num_to_letter(n: int) -> str:
    result = ""
    while n > 0:
        n -= 1
        result = chr(ord("A") + n % 26) + result
        n //= 26
    return result


def parse_meal_offsets(meals_str: str) -> set:
    offsets = set()
    for m in re.split(r"[,;\n]", meals_str or ""):
        m = m.strip().lower()
        if m in MEAL_MAP:
            offsets.add(MEAL_MAP[m])
    return offsets


def build_meal_updates(ws, date_start_str, date_end_str, pitanie_start, pitanie_end, total_people, row):
    if not date_start_str or not date_end_str or total_people <= 0:
        return []
    try:
        d_start = datetime.strptime(date_start_str.strip(), "%d.%m.%Y")
        d_end = datetime.strptime(date_end_str.strip(), "%d.%m.%Y")
    except ValueError:
        return []
    header3 = ws.row_values(3)
    date_pattern = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
    date_col_map = {}
    for i, val in enumerate(header3):
        val = str(val).strip()
        if date_pattern.match(val):
            date_col_map[val] = i + 1  # 1-indexed
    if not date_col_map:
        return []
    meals_arrival = parse_meal_offsets(pitanie_start)
    meals_departure = parse_meal_offsets(pitanie_end)
    standard = set(STANDARD_MEAL_OFFSETS)
    updates = []
    current = d_start
    while current <= d_end:
        date_key = current.strftime("%d.%m.%Y")
        if date_key in date_col_map:
            base_col = date_col_map[date_key]
            if current == d_start and current == d_end:
                active = meals_arrival | meals_departure if (meals_arrival or meals_departure) else standard
            elif current == d_start:
                active = meals_arrival if meals_arrival else standard
            elif current == d_end:
                active = meals_departure if meals_departure else standard
            else:
                active = standard
            for offset in sorted(active):
                col_letter = col_num_to_letter(base_col + offset)
                updates.append({"range": f"{col_letter}{row}", "values": [[total_people]]})
        current += timedelta(days=1)
    return updates


def get_creds():
    creds_json = os.environ.get("GOOGLE_CREDENTIALS_JSON")
    if creds_json:
        info = json.loads(creds_json, strict=False)
        return Credentials.from_service_account_info(info, scopes=SCOPES)
    return Credentials.from_service_account_file("/etc/secrets/credentials.json", scopes=SCOPES)


def get_worksheet(turnir: str) -> gspread.Worksheet:
    creds = get_creds()
    client = gspread.authorize(creds)
    return client.open_by_key(SPREADSHEET_ID).worksheet(turnir)


def find_first_empty_row(ws: gspread.Worksheet) -> int:
    col_b = ws.col_values(2)
    for row_idx in range(FIRST_DATA_ROW, len(col_b) + 2):
        val = col_b[row_idx - 1] if row_idx - 1 < len(col_b) else ""
        if val == "":
            return row_idx
    return max(len(col_b) + 1, FIRST_DATA_ROW)


def meal_label(meals_str: str) -> str:
    """Return comma-separated meal names from pitanie string."""
    parts = []
    for m in re.split(r"[,;\n]", meals_str or ""):
        m = m.strip().lower()
        if m in MEAL_MAP:
            parts.append(m)
    return ", ".join(parts)


def build_summary(
    name_team: str,
    turnir: str,
    name_zakazchik: str,
    phone: str,
    date_start: str,
    time_start: str,
    date_end: str,
    time_end: str,
    kol_detey: str,
    kol_trener: str,
    kol_parent: str,
    pitanie_start: str,
    pitanie_end: str,
) -> str:
    try:
        n_sportsmen = int(kol_detey or 0) + int(kol_parent or 0)
        n_trener = int(kol_trener or 0)
    except ValueError:
        n_sportsmen = 0
        n_trener = 0

    arrival = f"{date_start} {time_start}".strip()
    departure = f"{date_end} {time_end}".strip()

    meal_arrival_label = meal_label(pitanie_start)
    meal_departure_label = meal_label(pitanie_end)

    lines = [
        f"{name_team}, {name_zakazchik} {phone}".strip(", "),
        f"заезд {arrival}",
        f"выезд {departure}",
        f"{n_sportsmen} спортсменов, {n_trener} тренер",
    ]
    if meal_arrival_label:
        lines.append(f"питание в день заезда {meal_arrival_label}")
    if meal_departure_label:
        lines.append(f"питание в день выезда {meal_departure_label}")

    return "\n".join(lines)


def parse_person_list(text: str) -> list:
    """Parse textarea person list into rows: each line → [fio, dob, phone]."""
    rows = []
    for line in (text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        # Try to split on commas: FIO, DOB, phone
        parts = [p.strip() for p in line.split(",")]
        if len(parts) >= 3:
            rows.append([parts[0], parts[1], parts[2]])
        elif len(parts) == 2:
            rows.append([parts[0], parts[1], ""])
        else:
            rows.append([line, "", ""])
    return rows


def set_cell_text(cell, text: str):
    """Clear a table cell and set plain text."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


def fill_document(
    turnir: str,
    name_team: str,
    name_zakazchik: str,
    phone: str,
    spisok_detey: str,
    spisok_vzrosly: str,
) -> bytes:
    """Fill Word template and return bytes of the filled document."""
    template_bytes = base64.b64decode(TEMPLATE_B64)
    doc = Document(io.BytesIO(template_bytes))

    # --- Paragraph 0: title + team + trainer ---
    para0 = doc.paragraphs[0]
    runs = para0.runs
    # run[1] → tournament name
    if len(runs) > 1:
        runs[1].text = f"«{turnir}»"
    # run[3] → team name (bold)
    if len(runs) > 3:
        runs[3].text = f"«{name_team}»"
    # run[9] → trainer info
    if len(runs) > 9:
        runs[9].text = f"{name_zakazchik}, {phone}".strip(", ")

    # --- Table: fill athlete rows ---
    if doc.tables:
        table = doc.tables[0]
        # Collect all people: athletes first, then adults
        athletes = parse_person_list(spisok_detey)
        adults = parse_person_list(spisok_vzrosly)
        all_people = athletes + adults

        # Data rows start at index 1 (row 0 is header)
        data_rows = table.rows[1:]

        for i, row in enumerate(data_rows):
            cells = row.cells
            if i < len(all_people):
                person = all_people[i]
                set_cell_text(cells[0], str(i + 1))    # №
                set_cell_text(cells[1], person[0])      # ФИО
                set_cell_text(cells[2], person[1])      # Дата рождения
                set_cell_text(cells[3], person[2])      # Телефон
                set_cell_text(cells[4], "")             # Посадочное место
            else:
                # Clear unused rows
                for cell in cells:
                    set_cell_text(cell, "")

        # If we have more people than template rows, add new rows
        if len(all_people) > len(data_rows):
            template_data_row = data_rows[-1] if data_rows else None
            for i in range(len(data_rows), len(all_people)):
                person = all_people[i]
                # Add new row by copying last data row's XML
                if template_data_row is not None:
                    new_tr = copy.deepcopy(template_data_row._tr)
                    table._tbl.append(new_tr)
                    # Access the newly appended row
                    new_row = table.rows[-1]
                    cells = new_row.cells
                    set_cell_text(cells[0], str(i + 1))
                    set_cell_text(cells[1], person[0])
                    set_cell_text(cells[2], person[1])
                    set_cell_text(cells[3], person[2])
                    set_cell_text(cells[4], "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def upload_to_drive(doc_bytes: bytes, filename: str) -> str:
    """Upload document to Google Drive and return public download link."""
    creds = get_creds()
    service = build("drive", "v3", credentials=creds)

    file_metadata = {"name": filename, "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    media = MediaIoBaseUpload(io.BytesIO(doc_bytes), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    file = service.files().create(body=file_metadata, media_body=media, fields="id").execute()
    file_id = file.get("id")

    # Make public
    service.permissions().create(
        fileId=file_id,
        body={"type": "anyone", "role": "reader"},
    ).execute()

    # Return direct download link
    return f"https://drive.google.com/uc?export=download&id={file_id}"


@app.post("/webhook")
async def webhook(request: Request):
    try:
        content_type = request.headers.get("content-type", "")
        if "application/json" in content_type:
            data = await request.json()
        else:
            form = await request.form()
            data = dict(form)
        turnir = (data.get("turnir") or "").strip()
        if not turnir:
            return JSONResponse({"status": "ok", "ignored": True})
        name             = (data.get("name")             or "").strip()n        phone            = normalize_phone((data.get("phone") or "").strip())
        email            = (data.get("email")            or "").strip()
        name_team        = (data.get("name_team")        or "").strip()
        name_zakazchik   = (data.get("name_zakazchik")   or "").strip()
        format_oplaty    = (data.get("format_oplaty")    or "").strip()
        date_start       = (data.get("date_start")       or "").strip()
        time_start       = (data.get("time_start")       or "").strip()
        info_pribytie    = (data.get("info_pribytie")    or "").strip()
        info_otpravlenye = (data.get("info_otpravlenye") or "").strip()
        date_end         = (data.get("date_end")         or "").strip()
        time_end         = (data.get("time_end")         or "").strip()
        kol_detey        = (data.get("kol_detey")        or "").strip()
        kol_trener       = (data.get("kol_trener")       or "").strip()
        kol_parent       = (data.get("kol_parent")       or "").strip()
        transfer         = (data.get("transfer")         or "").strip()
        pitanie_start    = (data.get("pitanie_start")    or "").strip()
        pitanie_end      = (data.get("pitanie_end")      or "").strip()
        pitanie_syh_end  = (data.get("pitanie_syh_end")  or "").strip()
        spisok_detey     = (data.get("spisok_detey")     or "").strip()
        spisok_vzrosly   = (data.get("spisok_vzrosly")   or "").strip()

        team_contact   = f"{name_team}, {name}" if name_team and name else (name_team or name)
        arrival_dt     = f"{date_start} {time_start}".strip()
        departure_dt   = f"{date_end} {time_end}".strip()
        try:
            total_people = int(kol_detey or 0) + int(kol_trener or 0) + int(kol_parent or 0)
        except ValueError:
            total_people = 0
        transfer_cost = calculate_transfer_cost(total_people) if transfer.lower().startswith("да") else 0
        suh_paek_val = "да" if pitanie_syh_end.lower().startswith("да") else ("нет" if pitanie_syh_end else "")

        # Build summary text
        summary = build_summary(
            name_team=name_team,
            turnir=turnir,
            name_zakazchik=name_zakazchik or name,
            phone=phone,
            date_start=date_start,
            time_start=time_start,
            date_end=date_end,
            time_end=time_end,
            kol_detey=kol_detey,
            kol_trener=kol_trener,
            kol_parent=kol_parent,
            pitanie_start=pitanie_start,
            pitanie_end=pitanie_end,
        )

        # Generate and upload document
        doc_link = ""
        try:
            doc_bytes = fill_document(
                turnir=turnir,
                name_team=name_team,
                name_zakazchik=name_zakazchik or name,
                phone=phone,
                spisok_detey=spisok_detey,
                spisok_vzrosly=spisok_vzrosly,
            )
            doc_name = f"Список «{name_team}, {turnir}»"
            doc_link = upload_to_drive(doc_bytes, doc_name)
        except Exception as doc_exc:
            logger.exception("Document generation/upload error: %s", doc_exc)

        ws  = get_worksheet(turnir)
        row = find_first_empty_row(ws)
        serial = row - FIRST_DATA_ROW + 1
        updates = [
            {"range": f"A{row}",        "values": [[serial]]},
            {"range": f"B{row}",        "values": [[team_contact]]},
            {"range": f"C{row}",        "values": [[arrival_dt]]},
            {"range": f"D{row}",        "values": [[departure_dt]]},
            {"range": f"G{row}",        "values": [[kol_detey]]},
            {"range": f"H{row}",        "values": [[kol_trener]]},
            {"range": f"I{row}",        "values": [[kol_parent]]},
            {"range": f"BL{row}",       "values": [[suh_paek_val]]},
            {"range": f"BN{row}",       "values": [[transfer_cost]]},
            {"range": f"BQ{row}",       "values": [[format_oplaty]]},
            {"range": f"BT{row}",       "values": [[info_pribytie]]},
            {"range": f"BU{row}",       "values": [[info_otpravlenye]]},
            {"range": f"BV{row}",       "values": [[name]]},
            {"range": f"BW{row}",       "values": [[phone]]},
            {"range": f"BX{row}",       "values": [[email]]},
            {"range": f"{SUMMARY_COL}{row}", "values": [[summary]]},
            {"range": f"{LINK_COL}{row}",    "values": [[doc_link]]},
        ]
        meal_updates = build_meal_updates(ws, date_start, date_end, pitanie_start, pitanie_end, total_people, row)
        updates.extend(meal_updates)
        ws.batch_update(updates, value_input_option="USER_ENTERED")
        return JSONResponse({"status": "ok", "sheet": turnir, "row": row})
    except gspread.exceptions.WorksheetNotFound:
        return JSONResponse({"status": "error", "detail": f"Worksheet '{turnir}' not found"})
    except Exception as exc:
        logger.exception("Unexpected error: %s", exc)
        return JSONResponse({"status": "error", "detail": str(exc)})


@app.get("/")
def healthcheck():
    return {"status": "running"}
